"""
Template-based DPR Generator with Direct HTML to DOCX Conversion
This module generates DPR using HTML templates and directly converts them to DOCX format.
"""

import base64
import os
import tempfile
from datetime import date
from io import BytesIO
from urllib.parse import urlencode

import geopandas as gpd
import requests
from django.core.mail import EmailMessage
from django.core.mail.backends.smtp import EmailBackend
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from shapely.geometry import Point

# For direct HTML to DOCX conversion (optional, can be installed if needed)
try:
    from htmldocx import HtmlToDocx
except ImportError:
    HtmlToDocx = None

try:
    import pypandoc
except ImportError:
    pypandoc = None

from nrm_app.settings import (
    EMAIL_HOST,
    EMAIL_HOST_PASSWORD,
    EMAIL_HOST_USER,
    EMAIL_PORT,
    EMAIL_TIMEOUT,
    EMAIL_USE_SSL,
    GEOSERVER_URL,
)
from plans.models import Plan
from utilities.logger import setup_logger

from .models import ODK_settlement
from .utils import get_vector_layer_geoserver, sync_db_odk  # Added sync_db_odk import

logger = setup_logger(__name__)


def get_plan(plan_id):
    """Get plan object by ID"""
    try:
        return Plan.objects.get(plan_id=plan_id)
    except Plan.DoesNotExist:
        return None


def get_settlement_count_for_plan(planid):
    """Get total count of settlements for a plan"""
    return (
        ODK_settlement.objects.filter(plan_id=planid)
        .exclude(status_re="rejected")
        .count()
    )


def get_data_for_settlement(planid):
    """Get settlement data for a plan"""
    return ODK_settlement.objects.filter(plan_id=planid).exclude(status_re="rejected")


def get_settlement_coordinates_for_plan(planid):
    """Get coordinates for all settlements in a plan"""
    settlements = (
        ODK_settlement.objects.filter(plan_id=planid)
        .exclude(status_re="rejected")
        .values("settlement_name", "latitude", "longitude")
    )
    return [
        (
            settlement["settlement_name"],
            settlement["latitude"],
            settlement["longitude"],
        )
        for settlement in settlements
    ]


def get_mws_uid_for_settlement_gdf(mws_gdf, lat, lon):
    """Get MWS UID for a settlement based on coordinates"""
    settlement_point = Point(lon, lat)
    intersecting_mws = mws_gdf[mws_gdf.intersects(settlement_point)]

    if not intersecting_mws.empty:
        mws_uid = intersecting_mws.iloc[0]["uid"]
        return mws_uid
    else:
        return None


def generate_wms_map_url(plan, bbox=None, width=800, height=600):
    """
    Generate a WMS GetMap URL for the settlement layer
    """
    # Construct layer name
    district_name = str(plan.district.district_name).lower().replace(" ", "_")
    block_name = str(plan.block.block_name).lower().replace(" ", "_")
    layer_name = f"resources:settlement_{plan.plan_id}_{district_name}_{block_name}"

    # Default bbox if not provided
    if bbox is None:
        settlements = get_settlement_coordinates_for_plan(plan.plan_id)
        if settlements:
            lats = [s[1] for s in settlements]
            lons = [s[2] for s in settlements]

            padding = 0.01
            bbox = f"{min(lons) - padding},{min(lats) - padding},{max(lons) + padding},{max(lats) + padding}"
        else:
            bbox = "78.0,23.0,79.0,24.0"

    # WMS parameters
    params = {
        "SERVICE": "WMS",
        "VERSION": "1.1.0",
        "REQUEST": "GetMap",
        "LAYERS": layer_name,
        "BBOX": bbox,
        "WIDTH": width,
        "HEIGHT": height,
        "SRS": "EPSG:4326",
        "FORMAT": "image/png",
        "TRANSPARENT": "true",
    }

    wms_url = f"{GEOSERVER_URL}/wms?{urlencode(params)}"
    logger.info(f"Generated WMS URL: {wms_url}")
    return wms_url


def prepare_section_a_data(plan):
    """
    Prepare data for Section A: Team Details
    """
    return {
        "facilitator_name": plan.facilitator_name,
        "project_name": plan.plan.replace("_", " ").title(),
        "process_involved": "PRA, Gram Sabha, Transect Walk, GIS Mapping",
        "description": (
            "This section gives brief information about the Project Name, "
            "facilitator details responsible for the preparation of the Detailed "
            "Project Report (DPR). The process begins with Community Consultations, "
            "involving active engagement with community members "
            "to identify their needs and resources."
        ),
    }


def prepare_section_b_data(plan, total_settlements, mws_fortnight):
    """
    Prepare data for Section B: Brief of Village
    """
    mws_gdf = gpd.GeoDataFrame.from_features(mws_fortnight["features"])

    settlement_mws_ids = []
    settlement_coordinates = get_settlement_coordinates_for_plan(plan.plan_id)

    for settlement_name, latitude, longitude in settlement_coordinates:
        mws_uid = get_mws_uid_for_settlement_gdf(mws_gdf, latitude, longitude)
        if mws_uid:
            settlement_mws_ids.append((settlement_name, mws_uid))

    intersecting_mws_ids = "; ".join(
        [f"{name}: {mws_id}" for name, mws_id in settlement_mws_ids]
    )

    # Calculate the centroid
    if intersecting_mws_ids:
        intersecting_mws = mws_gdf[
            mws_gdf["uid"].isin(
                [mws_id.split(": ")[1] for mws_id in intersecting_mws_ids.split("; ")]
            )
        ]
        if not intersecting_mws.empty:
            centroid = intersecting_mws.geometry.unary_union.centroid
            latitude = f"{centroid.y:.8f}"
            longitude = f"{centroid.x:.8f}"
        else:
            latitude = "N/A"
            longitude = "N/A"
    else:
        latitude = "N/A"
        longitude = "N/A"

    return {
        "village_name": plan.village_name,
        "gram_panchayat": plan.gram_panchayat,
        "tehsil": plan.block.block_name,
        "district": plan.district.district_name,
        "state": plan.state.state_name,
        "total_settlements": str(total_settlements),
        "intersecting_mws_ids": intersecting_mws_ids or "N/A",
        "latitude": latitude,
        "longitude": longitude,
        "description": (
            "This section gives a brief overview of the village, "
            "including its name, associated Gram Panchayat, location "
            "details (block, district, and state), the number of settlements, "
            "intersecting micro watershed IDs, and the geographic coordinates "
            "(latitude and longitude) of the village."
        ),
    }


def fetch_wms_map_image(plan, for_docx=False):
    """
    Fetch map image from GeoServer WMS
    Returns either a base64 string (for HTML) or a file path (for DOCX)
    """
    try:
        wms_url = generate_wms_map_url(plan)

        # Fetch the image
        response = requests.get(wms_url, timeout=30)
        response.raise_for_status()

        if for_docx:
            # Save to temporary file for DOCX embedding
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp_file:
                tmp_file.write(response.content)
                tmp_path = tmp_file.name
            logger.info(f"Map image saved to temporary file: {tmp_path}")
            return tmp_path
        else:
            # Return as base64 for HTML embedding
            image_base64 = base64.b64encode(response.content).decode("utf-8")
            image_data_url = f"data:image/png;base64,{image_base64}"
            logger.info("Map image converted to base64 for HTML")
            return image_data_url

    except Exception as e:
        logger.error(f"Error fetching WMS map: {str(e)}")
        return None


def prepare_map_section_data(plan, mws_gdf=None, for_docx=False):
    """
    Prepare data for Map Section
    """
    # Get settlement data for map
    settlements = get_settlement_coordinates_for_plan(plan.plan_id)

    # Generate map image (as file path for DOCX, base64 for HTML)
    map_image = fetch_wms_map_image(plan, for_docx=for_docx)

    # Correct layer name format
    district_name = plan.district.district_name.lower()
    block_name = plan.block.block_name.lower()

    return {
        "title": "Settlement Map",
        "description": (
            "This map shows the geographical distribution of settlements within the project area. "
            "The settlement boundaries are fetched from the GeoServer resources workspace, "
            "displaying the spatial layout and relationships between different settlements."
        ),
        "map_image": map_image,  # Either file path or base64
        "settlements": settlements,
        "layer_info": {
            "workspace": "resources",
            "layer_name": f"settlement_{plan.plan_id}_{district_name}_{block_name}",
            "geoserver_url": GEOSERVER_URL.rstrip("/"),
        },
    }


def generate_dpr_html_for_docx(plan_id):
    """
    Generate DPR HTML optimized for DOCX conversion
    This version handles images properly for DOCX
    """
    logger.info("Generating HTML for DOCX conversion for plan ID: %s", plan_id)

    # Get plan
    plan = get_plan(plan_id)
    if not plan:
        logger.error("Plan not found for ID: %s", plan_id)
        return None, None

    # Sync ODK data
    logger.info("Syncing ODK data with database")
    sync_db_odk()

    # Get data
    total_settlements = get_settlement_count_for_plan(plan.plan_id)

    mws_fortnight = get_vector_layer_geoserver(
        geoserver_url=GEOSERVER_URL,
        workspace="mws_layers",
        layer_name="deltaG_fortnight_"
        + str(plan.district.district_name).lower().replace(" ", "_")
        + "_"
        + str(plan.block.block_name).lower().replace(" ", "_"),
    )

    # Get map data - with file path for DOCX
    map_data = prepare_map_section_data(plan, for_docx=True)

    # Prepare context
    context = {
        "title": "Detailed Project Report",
        "date": date.today().strftime("%B %d, %Y"),
        "plan_name": plan.plan,
        "section_a": prepare_section_a_data(plan),
        "section_b": prepare_section_b_data(plan, total_settlements, mws_fortnight),
        "map_data": map_data,
        "settlements": get_settlement_coordinates_for_plan(plan.plan_id),
    }

    # Return both HTML content and map data (with file path)
    return context, map_data


def create_dpr_document_template(plan):
    """
    Main function to create DPR document using direct HTML to DOCX conversion
    """
    logger.info("Creating DPR document for plan ID: %s", plan.plan_id)

    # Generate context and map data
    context, map_data = generate_dpr_html_for_docx(plan.plan_id)

    if not context:
        logger.error("Failed to generate context data")
        return None

    try:
        # Create a new document
        doc = Document()

        # Add title
        heading = doc.add_heading(context["title"], 0)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add date
        p = doc.add_paragraph(context["date"])
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add page break
        doc.add_page_break()

        # Section A: Team Details
        doc.add_heading("Section A: Team Details", level=1)
        doc.add_paragraph(context["section_a"]["description"])

        # Create team details table
        table = doc.add_table(rows=2, cols=3)
        table.style = "Table Grid"

        # Headers
        headers = [
            "Name of the Facilitator",
            "Project Name",
            "Process involved in the preparation of the DPR PRA",
        ]
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
            # Make header bold
            for paragraph in table.rows[0].cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Data
        table.rows[1].cells[0].text = context["section_a"]["facilitator_name"]
        table.rows[1].cells[1].text = context["section_a"]["project_name"]
        table.rows[1].cells[2].text = context["section_a"]["process_involved"]

        # Section B: Brief of Village
        doc.add_heading("Section B: Brief of Village", level=1)
        doc.add_paragraph(context["section_b"]["description"])

        # Create village brief table
        table_b = doc.add_table(rows=8, cols=2)
        table_b.style = "Table Grid"

        village_data = [
            ("Name of the Village", context["section_b"]["village_name"]),
            ("Name of the Gram Panchayat", context["section_b"]["gram_panchayat"]),
            ("Tehsil", context["section_b"]["tehsil"]),
            ("District", context["section_b"]["district"]),
            ("State", context["section_b"]["state"]),
            (
                "Number of Settlements in the Village",
                context["section_b"]["total_settlements"],
            ),
            (
                "Intersecting Micro Watershed IDs",
                context["section_b"]["intersecting_mws_ids"],
            ),
            (
                "Latitude and Longitude of the Village",
                f"{context['section_b']['latitude']}, {context['section_b']['longitude']}",
            ),
        ]

        for i, (label, value) in enumerate(village_data):
            table_b.rows[i].cells[0].text = label
            table_b.rows[i].cells[1].text = value
            # Make label bold
            for paragraph in table_b.rows[i].cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Map Section
        doc.add_heading("Settlement Map", level=1)
        doc.add_paragraph(map_data["description"])

        # Add map image if available
        if map_data.get("map_image") and os.path.exists(map_data["map_image"]):
            try:
                doc.add_picture(map_data["map_image"], width=Inches(6))

                # Add caption
                p = doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.add_run(
                    f"Figure: Settlement boundaries for {plan.plan}"
                ).italic = True

                # Clean up temporary file
                os.unlink(map_data["map_image"])
                logger.info("Map image added and temporary file cleaned up")

            except Exception as e:
                logger.error(f"Error adding map image: {str(e)}")
                doc.add_paragraph("Map could not be loaded.")

        # Add settlement coordinates table if available
        if context.get("settlements"):
            doc.add_heading("Settlement Coordinates", level=2)

            settlements_table = doc.add_table(
                rows=len(context["settlements"]) + 1, cols=3
            )
            settlements_table.style = "Table Grid"

            # Headers
            headers = ["Settlement Name", "Latitude", "Longitude"]
            for i, header in enumerate(headers):
                settlements_table.rows[0].cells[i].text = header
                for paragraph in settlements_table.rows[0].cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            # Data
            for i, settlement in enumerate(context["settlements"], start=1):
                settlements_table.rows[i].cells[0].text = settlement[0]
                settlements_table.rows[i].cells[1].text = f"{settlement[1]:.6f}"
                settlements_table.rows[i].cells[2].text = f"{settlement[2]:.6f}"

        logger.info("DPR document created successfully")
        return doc

    except Exception as e:
        logger.error(f"Error creating DPR document: {str(e)}")
        logger.exception(e)  # This will log the full traceback
        return None


def create_dpr_document_template_alternative(plan):
    """
    Alternative method using pypandoc for better conversion
    Requires: pip install pypandoc
    """
    import pypandoc

    logger.info("Creating DPR document using pypandoc for plan ID: %s", plan.plan_id)

    # Generate HTML content
    html_content, _ = generate_dpr_html_for_docx(plan.plan_id)

    if not html_content:
        logger.error("Failed to generate HTML content")
        return None

    # Save HTML to temporary file
    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".html", delete=False
    ) as tmp_html:
        tmp_html.write(html_content)
        html_path = tmp_html.name

    # Convert using pandoc
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
        docx_path = tmp_docx.name

    try:
        # Convert HTML to DOCX using pandoc
        pypandoc.convert_file(
            html_path,
            "docx",
            outputfile=docx_path,
            extra_args=["--standalone", "--self-contained"],
        )

        # Load the generated DOCX
        doc = Document(docx_path)

        # Clean up temp files
        os.unlink(html_path)
        os.unlink(docx_path)

        logger.info("DPR document created successfully using pypandoc")
        return doc

    except Exception as e:
        logger.error(f"Error converting with pypandoc: {str(e)}")
        # Fallback to htmldocx method
        return create_dpr_document_template(plan)


def send_dpr_email_template(doc, email_id, plan_name):
    """
    Send DPR document via email
    """
    try:
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        doc_bytes = buffer.getvalue()
        buffer.close()

        email_body = f"""
        Hi,
        Find attached the Detailed Project Report for {plan_name}.

        This DPR was generated using direct HTML to DOCX conversion.

        Best Regards,
        NRM Team
        """

        backend = EmailBackend(
            host=EMAIL_HOST,
            port=EMAIL_PORT,
            username=EMAIL_HOST_USER,
            password=EMAIL_HOST_PASSWORD,
            use_ssl=EMAIL_USE_SSL,
            timeout=EMAIL_TIMEOUT,
        )

        email = EmailMessage(
            subject=f"DPR of plan: {plan_name}",
            body=email_body,
            from_email=EMAIL_HOST_USER,
            to=[email_id],
            connection=backend,
        )

        email.attach(
            f"DPR_{plan_name}.docx",
            doc_bytes,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        logger.info("Sending DPR email to %s", email_id)
        email.send(fail_silently=False)
        logger.info("DPR email sent.")
        backend.close()

    except Exception as e:
        logger.exception("Error sending DPR email: %s", str(e))
        raise
